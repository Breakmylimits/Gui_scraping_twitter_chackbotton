from tkinter import * 
from tkinter import ttk
from PIL import ImageTk,Image
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from bs4 import BeautifulSoup as soup
import time
from docx import Document
from docx.shared import Inches
import urllib.request
import urllib

def send(): 
    id = text_find.get()
    save_pic = savech.get()
    hight = 0
    
    url= id
    driver = webdriver.Chrome()
    driver.get(url)
    
    time.sleep(5)

    document = Document()

    document.add_heading('POST TWITTER', 0)

    p = document.add_paragraph(' ')

    for i in range(10):  
        page_html = driver.page_source
        
        data = soup(page_html, 'html.parser')
    
        tweet = data.findAll('div',{'class': 'css-901oao r-18jsvk2 r-1qd0xha r-a023e6 r-16dba41 r-ad9z0x r-bcqeeo r-bnwqim r-qvutc0'})
        user = data.findAll('div',{'class':'css-901oao css-bfa6kz r-m0bqgq r-18u37iz r-1qd0xha r-a023e6 r-16dba41 r-ad9z0x r-bcqeeo r-qvutc0'})
        imge = data.findAll('img',{'class' : 'css-9pa8cd' })
        result = []   

        driver.execute_script("window.scrollTo(0,{})".format(hight))
        hight += 20000
        time.sleep(3) 

        for i,tw in enumerate(zip(user,tweet,imge)):
            i = i+1
            print(i)
            user =tw[0].text
            pic = tw[2]['src']             
            print('user : '+user)
            print('PIC : '+pic)
                
            comment = tw[1].text
            print('comment : '+comment)
            print('--------------------------------------------------')

            if save_pic==1:
                imgePic=open(user+".jpeg",'wb')
                imgePic.write(urllib.request.urlopen(pic).read())
                imgePic.close()
            else:
                pass

            p.add_run('ชื่อผู้ใช้ : '+user+'\n').bold = True
            p.add_run('comment : '+comment+'\n\n')         
            document.save('data Twitter.docx')

        
            
        
    driver.close()

    return     
     



Gui = Tk()
Gui.geometry('300x300')
Gui.title('Generate data from twitter')
Gui.iconbitmap('C:/Users/Nitro/Desktop/เรียนภาษาอังกฤษ/Gui_scraping_twitter-main/employee.ico')


FONT = ('Angsana New', 20)

Text1 = StringVar()
Text1.set('ใส่ Url twitter ที่ต้องการข้อมูล')
Label1 = ttk.Label(Gui, textvariable= Text1 , font=FONT ,width=40)
Label1.place(x=27,y=25)

text_find = StringVar()
bfind = ttk.Entry(Gui, textvariable= text_find , font=FONT ,width=20)
bfind.place(x=27,y=100)

Bdata = ttk.Button(Gui, text = 'start' ,command = send) 
Bdata.place(x=27,y=200)
savech = IntVar()
save = Checkbutton(Gui, text ='save Profile picture', variable= savech)
save.place(x = 27, y = 150) 

Gui.mainloop()